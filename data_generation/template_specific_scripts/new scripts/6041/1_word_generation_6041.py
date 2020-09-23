#!/usr/bin/env python
# coding: utf-8

# In[ ]:


import pandas as pd
df = pd.read_excel("1.xlsx")


# In[ ]:


############ PAPER 6041 FINAL  ####################

import re
from docx import Document

temp = ""
def docx_replace_regex(doc_obj, regex , replace, temp):
    for p in doc_obj.paragraphs:
        ##print(p.text)
        inline = p.runs
        for i in range(len(inline)):
            #print(inline[i].text)
  
        
        #print("####################")
        if regex in p.text:
            #inline = p.runs
            ##print(inline)
            # Loop added to work with runs (strings with same style)
            for i in range(len(inline)):  
                
                temp = temp + inline[i].text
                ##print(temp)
                if regex in temp:
                    if replace=="na":
                        ##print("!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!")
                        continue 
                    else:
                        temp = temp.replace(regex, replace)
                        ##print(temp)
                        inline[i].text=temp        
        
#    for i in range(len(inline)):
#        #print(inline[i].text)
  
def printf (doc_obj):
    
    for p in doc_obj.paragraphs:
        if p.text=='':
            continue
        else:
            #print(p.text)
            ##print(p.text)
       
        #try:
            ##print(inline[0].text)
        #except:
        #    continue
    #print(" \n ############################################################################# \n")

    
    
    

text_to_replace1 = "dddddddddd"
text_to_replace2 = "dfgdfgdgggd"
#text_to_replace3 = "ÖZP, 32 (2003) 4"
text_to_replace4 = "dfdfdfdfdfdfdfdsdsdsds"
#text_to_replace5 = "ZeS-Arbeitspapier Nr. 07/2004"
#text_to_replace6 = "TUTS-WP-1-2005"


count=0
number=0
stopnumber=0
for row in df.index:

    replace1 = df['author'][count+1000]
    replace2 = df['title'][count+1000]
    #replace3 = df['journal'][count]
    replace4 = df['abstract'][count+1000].replace('\xa0',' ').replace('                ',' ').replace("  "," ").replace("  "," ").replace("  "," ").replace("  "," ").replace("  "," ").replace("  "," ").replace("  "," ")
    #replace5 = df['journal'][count]
    #replace5 = ""
    #replace6 = ""
    ##print(replace1)
    count=count+1
    
    
    
    filename = "6041.docx"
    doc = Document(filename)
    docx_replace_regex(doc, text_to_replace1 , replace1, temp)
    docx_replace_regex(doc, text_to_replace2 , replace2, temp)
    #docx_replace_regex(doc, text_to_replace3 , replace3, temp)
    docx_replace_regex(doc, text_to_replace4 , replace4, temp)
    #docx_replace_regex(doc, text_to_replace5 , replace5, temp)
    #docx_replace_regex(doc, text_to_replace6 , replace6, temp)
    
    
    for section in doc.sections:
        footer = section.footer
        ##print("ssssssssssssssssssssssssssssssssssssssssssssssssssssssssssssss")
        #print(footer.paragraphs[0].text)
        ##print("ssssssssssssssssssssssssssssssssssssssssssssssssssssssssssssss")
        #footer.paragraphs[0].text  = footer.paragraphs[0].text.replace("jjj", "fredfgfdgdfgfdgfdgdfgdfgdf")
        footer.paragraphs[1].text  = df['journal'][count+1000]+' ( '+str(df['year'][count+1000])+' )'
    
    printf (doc)
    
    #print(" \n ---------------------------------------------------------------------------------------- \n")

    doc.save('6041_' + str(number) +  '.docx')
    number = number+1
    stopnumber=stopnumber+1
    if stopnumber==1600:
        break

    
#regex1 = re.compile(r"Der vorliegende Bericht befaßt sich exemplarisch mit Techniken der Machteroberung und Machtbehauptung der russischen/sowjetischen Kommunisten seit der Oktoberrevolution von 1917 sowie mit Bestrebungen, aus der kommunistischen Bewegung heraus reformorientierte Gegenbewegungen zu bilden. Dabei wird an einigen Punkten versucht, Querverbindungen zu der aus Ruinen auferstandenen KP Rußlands unter Sjuganow zu ziehen. Die Analyse stützt sich auf Originalquellen der betroffenen Parteien und zieht darüber hinaus Beiträge vornehm-lich russischer und deutscher Spezialisten heran.")
#replace1 = "AAAAAAAAAAAAAAA"


# filename = "1181.docx"
# doc = Document(filename)
# docx_replace_regex(doc, text_to_replace1 , replace1, temp)
# docx_replace_regex(doc, text_to_replace2 , replace2, temp)
# docx_replace_regex(doc, text_to_replace3 , replace3, temp)
# docx_replace_regex(doc, text_to_replace4 , replace4, temp)
# docx_replace_regex(doc, text_to_replace5 , replace5, temp)
# docx_replace_regex(doc, text_to_replace6 , replace6, temp)

# doc.save('1181e.docx')

